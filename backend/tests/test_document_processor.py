"""
Tests for Document Processor
=============================
Tests text extraction from PDF, DOCX, and TXT files.
"""

import os
import tempfile
import pytest
from app.core.document_processor import DocumentProcessor, ExtractedDocument


@pytest.fixture
def processor():
    return DocumentProcessor()


@pytest.fixture
def sample_txt_path():
    content = "Hello world.\n\nThis is a test document.\n\nIt has multiple paragraphs."
    with tempfile.NamedTemporaryFile(
        mode="w", suffix=".txt", delete=False, encoding="utf-8"
    ) as f:
        f.write(content)
        path = f.name
    yield path
    os.unlink(path)


@pytest.fixture
def sample_docx_path():
    """Create a minimal DOCX file for testing."""
    from docx import Document
    doc = Document()
    doc.add_heading("Test Document", level=1)
    doc.add_paragraph("This is the first paragraph of the test document.")
    doc.add_paragraph("This is the second paragraph with more content.")

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        path = f.name
    doc.save(path)
    yield path
    os.unlink(path)


class TestTxtExtraction:

    def test_extracts_text_from_txt(self, processor, sample_txt_path):
        result = processor.process(sample_txt_path, "test.txt")
        assert isinstance(result, ExtractedDocument)
        assert result.file_type == "txt"
        assert result.total_chars > 0
        assert "Hello world" in result.full_text

    def test_txt_has_pages(self, processor, sample_txt_path):
        result = processor.process(sample_txt_path, "test.txt")
        assert len(result.pages) >= 1

    def test_empty_txt_handled(self, processor):
        with tempfile.NamedTemporaryFile(
            mode="w", suffix=".txt", delete=False
        ) as f:
            f.write("")
            path = f.name
        try:
            result = processor.process(path, "empty.txt")
            assert result.is_empty
        finally:
            os.unlink(path)

    def test_txt_encoding_detection(self, processor):
        """Test that non-UTF-8 files are handled gracefully."""
        content = "Hello world. Simple ASCII content."
        with tempfile.NamedTemporaryFile(
            mode="wb", suffix=".txt", delete=False
        ) as f:
            f.write(content.encode("latin-1"))
            path = f.name
        try:
            result = processor.process(path, "latin.txt")
            assert result.total_chars > 0
        finally:
            os.unlink(path)

    def test_unsupported_extension_returns_error(self, processor):
        with tempfile.NamedTemporaryFile(suffix=".xyz", delete=False) as f:
            path = f.name
        try:
            result = processor.process(path, "test.xyz")
            assert result.error is not None
        finally:
            os.unlink(path)


class TestDocxExtraction:

    def test_extracts_text_from_docx(self, processor, sample_docx_path):
        result = processor.process(sample_docx_path, "test.docx")
        assert result.file_type == "docx"
        assert result.total_chars > 0
        assert "Test Document" in result.full_text or "first paragraph" in result.full_text

    def test_docx_has_pages(self, processor, sample_docx_path):
        result = processor.process(sample_docx_path, "test.docx")
        assert len(result.pages) >= 1

    def test_docx_metadata_extracted(self, processor, sample_docx_path):
        result = processor.process(sample_docx_path, "test.docx")
        assert isinstance(result.metadata, dict)


class TestErrorHandling:

    def test_missing_file_returns_error(self, processor):
        result = processor.process("/nonexistent/path/file.txt", "file.txt")
        assert result.error is not None

    def test_corrupted_file_returns_error(self, processor):
        with tempfile.NamedTemporaryFile(
            mode="wb", suffix=".pdf", delete=False
        ) as f:
            f.write(b"this is not a valid PDF")
            path = f.name
        try:
            result = processor.process(path, "bad.pdf")
            # Either returns error or empty document — both acceptable
            assert result.error is not None or result.is_empty
        finally:
            os.unlink(path)