"""
Document Processor
==================
Extracts raw text from PDF, DOCX, and TXT files.

Why separate this from chunking?
- Single Responsibility: this file ONLY extracts text
- Testable: we can test PDF extraction without testing chunking
- Swappable: if we want to add OCR later, we only change this file

Library choices:
- PyMuPDF (fitz): Best PDF library — handles complex layouts,
  extracts page numbers, faster than pypdf
- python-docx: Official DOCX library by Microsoft
- chardet: Detects file encoding (UTF-8, Latin-1, etc.)
  so we don't crash on non-English documents
"""

import logging
import os
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

import chardet
import fitz  # PyMuPDF

logger = logging.getLogger(__name__)


@dataclass
class ExtractedPage:
    """
    Represents one page of extracted content.
    Keeping pages separate lets us report accurate page numbers in citations.
    """
    page_number: int        # 1-indexed (human readable)
    text: str               # Raw text from this page
    char_count: int = 0

    def __post_init__(self):
        self.char_count = len(self.text)


@dataclass
class ExtractedDocument:
    """
    Complete extraction result for one document.
    Contains all pages plus document-level metadata.
    """
    filename: str
    file_type: str                          # "pdf", "docx", "txt"
    pages: list[ExtractedPage]
    total_pages: int
    total_chars: int
    metadata: dict = field(default_factory=dict)
    error: Optional[str] = None

    @property
    def full_text(self) -> str:
        """All pages joined — used for simple text access."""
        return "\n\n".join(p.text for p in self.pages if p.text.strip())

    @property
    def is_empty(self) -> bool:
        return self.total_chars < 50  # Less than 50 chars = effectively empty


class DocumentProcessor:
    """
    Orchestrates document text extraction.

    Usage:
        processor = DocumentProcessor()
        doc = processor.process(file_path="/uploads/report.pdf", filename="report.pdf")
        print(doc.total_pages)   # 42
        print(doc.pages[0].text) # "Chapter 1: Introduction..."
    """

    def process(self, file_path: str, filename: str) -> ExtractedDocument:
        """
        Main entry point. Detects file type and delegates to the
        appropriate extractor.

        Args:
            file_path: Absolute path to the file on disk
            filename:  Original filename (used to detect type + for metadata)

        Returns:
            ExtractedDocument with all pages and metadata
        """
        ext = Path(filename).suffix.lower().lstrip(".")
        logger.info(f"Processing {filename} ({ext}) from {file_path}")

        try:
            if ext == "pdf":
                return self._extract_pdf(file_path, filename)
            elif ext == "docx":
                return self._extract_docx(file_path, filename)
            elif ext == "txt":
                return self._extract_txt(file_path, filename)
            else:
                raise ValueError(f"Unsupported file type: .{ext}")
        except Exception as e:
            logger.exception(f"Failed to process {filename}")
            return ExtractedDocument(
                filename=filename,
                file_type=ext,
                pages=[],
                total_pages=0,
                total_chars=0,
                error=str(e),
            )

    # ── PDF Extraction ───────────────────────────────────────

    def _extract_pdf(self, file_path: str, filename: str) -> ExtractedDocument:
        """
        Extract text from PDF using PyMuPDF (fitz).

        PyMuPDF advantages over pypdf:
        - 5-10x faster
        - Better handling of complex layouts (columns, tables)
        - Preserves reading order better
        - Handles encrypted PDFs (with empty password)
        - Gives us page-level extraction natively

        The main challenge with PDFs:
        - Some PDFs are "image PDFs" — scanned documents where
          the text is actually a picture. PyMuPDF returns empty
          strings for these pages. Real solution = OCR (Tesseract).
          For now we warn about empty pages.
        """
        doc = fitz.open(file_path)
        pages = []
        total_chars = 0

        # Extract document-level metadata
        metadata = {}
        pdf_meta = doc.metadata
        if pdf_meta:
            metadata = {
                "title": pdf_meta.get("title", ""),
                "author": pdf_meta.get("author", ""),
                "subject": pdf_meta.get("subject", ""),
                "creator": pdf_meta.get("creator", ""),
            }
            # Remove empty values
            metadata = {k: v for k, v in metadata.items() if v}

        for page_num in range(len(doc)):
            page = doc[page_num]

            # extract_text() returns plain text.
            # "blocks" sort: orders text blocks by reading position
            # This gives better results for multi-column layouts
            text = page.get_text("text", sort=True)

            # Clean the extracted text
            text = self._clean_pdf_text(text)

            if not text.strip():
                logger.warning(
                    f"Page {page_num + 1} of {filename} is empty — "
                    "may be a scanned image page"
                )

            extracted_page = ExtractedPage(
                page_number=page_num + 1,  # Convert to 1-indexed
                text=text,
            )
            pages.append(extracted_page)
            total_chars += extracted_page.char_count

        doc.close()

        logger.info(
            f"PDF extracted: {filename}, "
            f"{len(pages)} pages, {total_chars:,} chars"
        )

        return ExtractedDocument(
            filename=filename,
            file_type="pdf",
            pages=pages,
            total_pages=len(pages),
            total_chars=total_chars,
            metadata=metadata,
        )

    def _clean_pdf_text(self, text: str) -> str:
        """
        Clean raw PDF text extraction artifacts.

        Common PDF extraction issues:
        - Ligatures rendered as special chars (ﬁ instead of fi)
        - Hyphenated words split across lines: "infor-\nmation"
        - Multiple spaces from column layouts
        - Page headers/footers mixed into body text
        """
        import re

        # Fix hyphenated line breaks (common in justified text)
        # "infor-\nmation" → "information"
        text = re.sub(r"-\n(\w)", r"\1", text)

        # Normalise whitespace
        text = re.sub(r" {2,}", " ", text)

        # Collapse more than 2 newlines
        text = re.sub(r"\n{3,}", "\n\n", text)

        return text.strip()

    # ── DOCX Extraction ──────────────────────────────────────

    def _extract_docx(self, file_path: str, filename: str) -> ExtractedDocument:
        """
        Extract text from DOCX using python-docx.

        DOCX structure:
        - Document → Paragraphs → Runs
        - Each paragraph is one block of text (heading, body, list item)
        - Runs are styled segments within a paragraph

        We treat each paragraph as part of a "page" — DOCX doesn't have
        real page boundaries (pages are computed by Word's renderer).
        We simulate pages by grouping every 50 paragraphs together.
        This gives reasonable page citations even for DOCX files.
        """
        from docx import Document
        from docx.oxml.ns import qn

        doc = Document(file_path)
        all_paragraphs = []

        # Extract core properties metadata
        metadata = {}
        try:
            core_props = doc.core_properties
            if core_props.title:
                metadata["title"] = core_props.title
            if core_props.author:
                metadata["author"] = core_props.author
            if core_props.subject:
                metadata["subject"] = core_props.subject
        except Exception:
            pass

        # Extract text from all paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Include heading level for context
                if para.style.name.startswith("Heading"):
                    level = para.style.name.replace("Heading ", "")
                    text = f"{'#' * int(level) if level.isdigit() else '#'} {text}"
                all_paragraphs.append(text)

        # Extract text from tables (often missed by naive extractors)
        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_texts.append(cell_text)
                if row_texts:
                    all_paragraphs.append(" | ".join(row_texts))

        # Group paragraphs into simulated "pages"
        paragraphs_per_page = 50
        pages = []
        total_chars = 0

        for i in range(0, max(len(all_paragraphs), 1), paragraphs_per_page):
            page_paras = all_paragraphs[i : i + paragraphs_per_page]
            page_text = "\n\n".join(page_paras)
            page_num = (i // paragraphs_per_page) + 1

            extracted_page = ExtractedPage(
                page_number=page_num,
                text=page_text,
            )
            pages.append(extracted_page)
            total_chars += extracted_page.char_count

        if not pages:
            pages = [ExtractedPage(page_number=1, text="")]

        logger.info(
            f"DOCX extracted: {filename}, "
            f"{len(doc.paragraphs)} paragraphs → {len(pages)} pages, "
            f"{total_chars:,} chars"
        )

        return ExtractedDocument(
            filename=filename,
            file_type="docx",
            pages=pages,
            total_pages=len(pages),
            total_chars=total_chars,
            metadata=metadata,
        )

    # ── TXT Extraction ───────────────────────────────────────

    def _extract_txt(self, file_path: str, filename: str) -> ExtractedDocument:
        """
        Extract text from plain text files.

        The main challenge with TXT files is encoding.
        Files can be UTF-8, Latin-1, Windows-1252, etc.
        chardet detects the encoding automatically.

        We split TXT into "pages" by character count
        to maintain consistent citation granularity.
        """
        # Detect encoding
        with open(file_path, "rb") as f:
            raw_bytes = f.read()

        detection = chardet.detect(raw_bytes)
        encoding = detection.get("encoding") or "utf-8"
        confidence = detection.get("confidence", 0)

        logger.info(
            f"TXT encoding detected: {encoding} "
            f"(confidence: {confidence:.0%})"
        )

        # Decode with detected encoding, fallback to UTF-8 with error replacement
        try:
            text = raw_bytes.decode(encoding)
        except (UnicodeDecodeError, LookupError):
            logger.warning(
                f"Failed to decode {filename} as {encoding}, "
                "falling back to UTF-8 with replacement"
            )
            text = raw_bytes.decode("utf-8", errors="replace")

        # Split into pages by character count
        chars_per_page = 3000
        pages = []
        total_chars = len(text)

        for i in range(0, max(len(text), 1), chars_per_page):
            page_text = text[i : i + chars_per_page]
            page_num = (i // chars_per_page) + 1

            extracted_page = ExtractedPage(
                page_number=page_num,
                text=page_text,
            )
            pages.append(extracted_page)

        if not pages:
            pages = [ExtractedPage(page_number=1, text="")]

        logger.info(
            f"TXT extracted: {filename}, "
            f"{total_chars:,} chars → {len(pages)} pages"
        )

        return ExtractedDocument(
            filename=filename,
            file_type="txt",
            pages=pages,
            total_pages=len(pages),
            total_chars=total_chars,
            metadata={"encoding": encoding},
        )